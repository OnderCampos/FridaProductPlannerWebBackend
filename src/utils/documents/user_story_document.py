import json
import base64
import logging
import os
import re
import secrets
import string
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import quote

from fastapi import HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from src.schemas.response import ResponseModel
from src.schemas.user_data import UserData
from src.intelligence.runtime import AgentName, run_agent
from src.services.setup.firebase_setup import FIREBASE, FIRESTORE_CLIENT
from src.services.setup.language_setup import get_default_llm_language, normalize_language
from src.utils.core.validation_utils import get_code_block
from src.utils.planning.epics import get_epic_by_id
from src.utils.planning.projects import get_project_by_id
from src.utils.planning.user_stories import get_user_story_by_id

logger = logging.getLogger(__name__)

USER_STORY_DOCUMENT_DRAFTS_COLLECTION = "user_story_document_drafts"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

MAX_DOCUMENT_QUESTIONS = 20
MAX_WIREFRAME_IMAGES = 10
MAX_WIREFRAME_IMAGE_BYTES = 5 * 1024 * 1024  # 5 MB each
MAX_INLINE_WIREFRAME_IMAGES = 3
MAX_INLINE_WIREFRAME_IMAGE_BYTES = 200 * 1024  # 200 KB each (Firestore draft storage safety)

WIREFRAME_IMAGES_KEY = "wireframe_mockup_images"

BRAND_NAVY_HEX = "0D2A4A"
BRAND_NAVY_DARK_HEX = "0A1E34"
BRAND_LIGHT_HEX = "F3F6FA"

DEFAULT_HEADER_IMAGE = Path(__file__).resolve().parent / "assets" / "header.png"
FALLBACK_HEADER_IMAGE = Path(__file__).resolve().parent / "assets" / "report_header.png"

DOCUMENT_SECTION_KEYS = [
    "description_and_scope",
    "out_of_scope",
    "preconditions",
    "entry_points",
    "output_points",
    "success_flow",
    "wireframe_mockup",
    "field_description",
    "api_description",
    "acceptance_criteria",
    "test_scenarios",
    "dependencies",
    "benefits",
    "estimation_dev",
]

DOCX_FORMALIZE_SECTION_KEYS = [
    "description_and_scope",
    "out_of_scope",
    "preconditions",
    "entry_points",
    "output_points",
    "success_flow",
    "field_description",
    "api_description",
    "acceptance_criteria",
    "test_scenarios",
    "benefits",
    "dependencies",
]

QUESTION_TEXT_BY_KEY = {
    "description_and_scope": "Description and Scope: Describe what this user story changes, where it applies, and what is included.",
    "out_of_scope": "Out of Scope: List what is explicitly NOT included in this user story.",
    "preconditions": "Preconditions: List any required setup, roles/permissions, data, or prior steps needed before this feature can be used.",
    "entry_points": "Entry Points: Describe the entry points/inputs needed to use this feature (screens, actions, inputs).",
    "output_points": "Output Points: Describe the outputs produced (screens, fields, messages, system effects).",
    "success_flow": "Success Flow: Describe the happy path step-by-step from entry to successful completion.",
    "wireframe_mockup": "Wireframe / Mockup: Upload one or more images (PNG/JPG) and/or add a short note. If not available, write N/A.",
    "field_description": "Field Description: List the key fields involved and business rules, or write N/A.",
    "api_description": "API Description: Describe any API/system interactions (source/target/format), or write N/A.",
    "acceptance_criteria": "Acceptance Criteria: Provide bullet points describing the conditions for the story to be considered done.",
    "test_scenarios": "Test Scenarios: List key test scenarios (happy path, edge cases, error cases).",
    "dependencies": "Dependencies: List any dependencies, prerequisites, or related stories/systems this user story relies on, or write N/A.",
    "benefits": "Benefits: Explain the expected benefits of this user story, or write N/A.",
    "estimation_dev": "Estimation Dev: Provide the time estimate (and any notes), or write N/A.",
}


QUESTION_TEXT_BY_KEY_ES = {
    "description_and_scope": "Descripción y alcance: Describe qué cambia esta historia de usuario, dónde aplica y qué está incluido.",
    "out_of_scope": "Fuera de alcance: Enumera lo que explícitamente NO está incluido en esta historia de usuario.",
    "preconditions": "Precondiciones: Enumera cualquier configuración requerida, roles/permisos, datos o pasos previos necesarios antes de usar esta funcionalidad.",
    "entry_points": "Puntos de entrada: Describe los puntos de entrada/insumos necesarios para usar esta funcionalidad (pantallas, acciones, entradas).",
    "output_points": "Puntos de salida: Describe las salidas producidas (pantallas, campos, mensajes, efectos en el sistema).",
    "success_flow": "Flujo de éxito: Describe el camino feliz paso a paso desde la entrada hasta la finalización exitosa.",
    "wireframe_mockup": "Wireframe / mockup: Sube una o más imágenes (PNG/JPG) y/o agrega una nota breve. Si no está disponible, escribe N/A.",
    "field_description": "Descripción de campos: Lista los campos clave involucrados y reglas de negocio, o escribe N/A.",
    "api_description": "Descripción de API: Describe cualquier interacción con APIs/sistemas (origen/destino/formato), o escribe N/A.",
    "acceptance_criteria": "Criterios de aceptación: Proporciona viñetas con las condiciones para considerar la historia como terminada.",
    "test_scenarios": "Escenarios de prueba: Lista los escenarios clave (camino feliz, casos borde y casos de error).",
    "benefits": "Beneficios: Explica los beneficios esperados de esta historia de usuario, o escribe N/A.",
    "estimation_dev": "Estimación (Dev): Proporciona la estimación de tiempo (y notas), o escribe N/A.",
}


def _get_question_text(key: str) -> str:
    language = normalize_language(get_default_llm_language(), default="English")
    if language == "Spanish":
        return QUESTION_TEXT_BY_KEY_ES.get(key, QUESTION_TEXT_BY_KEY.get(key, key))
    return QUESTION_TEXT_BY_KEY.get(key, key)


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _now_display_date() -> str:
    # Template uses DD.MM.YYYY / DD/MM/YYYY in examples.
    return datetime.now(timezone.utc).strftime("%d.%m.%Y")


def _random_suffix(length: int = 10) -> str:
    alphabet = string.ascii_uppercase + string.digits
    return "".join(secrets.choice(alphabet) for _ in range(max(1, length)))


def _is_blank(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        normalized = value.strip()
        if not normalized:
            return True
        if normalized.lower() in {"n/a", "na", "none", "null", "-"}:
            return True
    return False


def _docx_formalization_enabled() -> bool:
    value = str(os.getenv("ENABLE_DOCX_FORMALIZATION") or "").strip().lower()
    # Default: enabled (can be disabled with 0/false/no/off).
    if not value:
        return True
    return value not in {"0", "false", "no", "off"}


def _safe_json_load(raw: str) -> Optional[Any]:
    text = str(raw or "").strip()
    if not text:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        try:
            block = get_code_block(text)
            if not block:
                return None
            return json.loads(block)
        except Exception:
            return None


async def _formalize_docx_document_fields(
    *,
    user_data: UserData,
    document: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Rewrite free-text user-provided content into a more formal, grammatically correct style
    without changing intent. Falls back to the original content on any error.
    """
    if not _docx_formalization_enabled():
        return document

    fields: Dict[str, str] = {}
    for key in DOCX_FORMALIZE_SECTION_KEYS:
        value = document.get(key)
        if _is_blank(value):
            continue
        text = str(value or "").strip()
        if not text:
            continue
        fields[key] = text

    if not fields:
        return document

    prompt = f"""
You are an expert technical writer.

Task:
Rewrite each field value to be more formal and grammatically correct while preserving the original meaning and intent.

Rules:
- Do NOT add new requirements, assumptions, or details.
- Do NOT remove important details.
- Keep the original language of each field (do not translate).
- Preserve formatting: keep bullet points, numbered steps, and line breaks when present.
- Return ONLY valid JSON with the exact same keys as the input.

Input JSON:
{json.dumps(fields, ensure_ascii=False, indent=2)}
""".strip()

    try:
        raw = await run_agent(AgentName.USER_STORY_DOCUMENT, prompt, user_data, model_tier="mini")
        parsed = _safe_json_load(raw)
        if not isinstance(parsed, dict):
            return document

        updated = dict(document)
        for key, original in fields.items():
            candidate = parsed.get(key)
            if not isinstance(candidate, str):
                continue
            cleaned = candidate.strip()
            if not cleaned:
                continue
            # Guardrail: avoid suspiciously short outputs that could drop content.
            if len(cleaned) < max(8, int(len(original) * 0.4)):
                continue
            updated[key] = cleaned
        return updated
    except Exception as exc:
        logger.warning("DOCX formalization skipped due to error: %s", exc)
        return document


def _resolve_header_image_path() -> Optional[Path]:
    """
    Resolve a header image to embed into the DOCX.

    Priority:
    1) env var USER_STORY_DOC_HEADER_IMAGE
    2) default repo asset `src/utils/documents/assets/header.png`
    """
    raw = str(os.getenv("USER_STORY_DOC_HEADER_IMAGE") or "").strip()
    if raw:
        candidate = Path(raw)
        if candidate.exists() and candidate.is_file():
            return candidate
    if DEFAULT_HEADER_IMAGE.exists() and DEFAULT_HEADER_IMAGE.is_file():
        return DEFAULT_HEADER_IMAGE
    if FALLBACK_HEADER_IMAGE.exists() and FALLBACK_HEADER_IMAGE.is_file():
        return FALLBACK_HEADER_IMAGE
    return None


def _apply_document_defaults(doc) -> None:
    try:
        from docx.shared import Inches  # type: ignore
    except Exception:
        Inches = None  # type: ignore

    if Inches is None:
        return

    _set_document_font(doc, "Arial")

    for section in doc.sections:
        # Leave enough room for the branded header image so the body starts
        # right below it, while still allowing the header itself to start at
        # the very top of the page (header distance is configured separately).
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def _set_document_font(doc, font_name: str) -> None:
    """
    Apply a default font for the whole document by updating common paragraph styles.
    """
    name = str(font_name or "").strip()
    if not name:
        return

    try:
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        qn = None  # type: ignore

    for style_name in ("Normal", "Header", "Footer"):
        try:
            style = doc.styles[style_name]
        except Exception:
            continue
        try:
            style.font.name = name
        except Exception:
            pass
        if qn is None:
            continue
        try:
            r_pr = style.element.get_or_add_rPr()
            r_fonts = r_pr.get_or_add_rFonts()
            r_fonts.set(qn("w:ascii"), name)
            r_fonts.set(qn("w:hAnsi"), name)
            r_fonts.set(qn("w:cs"), name)
            r_fonts.set(qn("w:eastAsia"), name)
        except Exception:
            continue


def _add_header_image(doc) -> bool:
    path = _resolve_header_image_path()
    if not path:
        return False

    try:
        from docx.shared import Inches  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except Exception:
        return False

    sections = []
    try:
        sections = list(doc.sections)
    except Exception:
        sections = []
    if not sections:
        return False

    inserted_any = False
    for section in sections:
        try:
            section.different_first_page_header_footer = False
        except Exception:
            pass
        try:
            section.different_odd_even_pages_header_footer = False
        except Exception:
            pass
        try:
            # Place the header at the very top of the page.
            section.header_distance = Inches(0)
        except Exception:
            pass

        try:
            full_width = section.page_width
        except Exception:
            full_width = Inches(8.5)

        headers = []
        try:
            # Use the default header only. Accessing `first_page_header` /
            # `even_page_header` can create extra header parts and confuse
            # some DOCX renderers when "different first page" / "odd-even"
            # are disabled.
            headers.append(section.header)
        except Exception:
            headers = []
        try:
            section.header.is_linked_to_previous = False
        except Exception:
            pass

        for header in headers:
            try:
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.text = ""
                try:
                    paragraph.paragraph_format.space_before = 0
                    paragraph.paragraph_format.space_after = 0
                    # Make the image full-bleed across the page by canceling the
                    # section's left/right margins inside the header paragraph.
                    paragraph.paragraph_format.left_indent = -section.left_margin
                    paragraph.paragraph_format.right_indent = -section.right_margin
                except Exception:
                    pass

                run = paragraph.add_run()
                run.add_picture(str(path), width=full_width)
                inserted_any = True
            except Exception:
                continue
    return inserted_any


def _add_report_header(doc) -> None:
    """
    Add a "report header" at the top of the document body so it is always visible
    (some viewers hide the Word header area by default).
    """
    path = _resolve_header_image_path()

    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except Exception:
        WD_ALIGN_PARAGRAPH = None  # type: ignore

    try:
        section = doc.sections[0]
        full_width = section.page_width
    except Exception:
        section = None
        full_width = None

    if path and full_width is not None:
        paragraph = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 10
            if section is not None:
                paragraph.paragraph_format.left_indent = -section.left_margin
                paragraph.paragraph_format.right_indent = -section.right_margin
        except Exception:
            pass
        run = paragraph.add_run()
        try:
            run.add_picture(str(path), width=full_width)
            return
        except Exception:
            # fall through to simple banner
            pass

    # Fallback banner: navy background with white text.
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("REPORT")
    run.bold = True
    _set_paragraph_shading(paragraph, BRAND_NAVY_HEX)
    _style_paragraph_text(paragraph, bold=True, color_hex="FFFFFF")
    try:
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 10
    except Exception:
        pass


def _set_paragraph_shading(paragraph, fill_hex: str) -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        return

    fill = (fill_hex or "").strip().lstrip("#").upper()
    if not fill:
        return

    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _style_paragraph_text(paragraph, *, bold: bool, color_hex: str) -> None:
    try:
        from docx.shared import RGBColor  # type: ignore
    except Exception:
        return

    color = (color_hex or "").strip().lstrip("#").upper()
    if not color:
        return
    for run in paragraph.runs:
        run.bold = bold
        try:
            run.font.color.rgb = RGBColor.from_string(color)
        except Exception:
            pass


def _add_section_header(doc, title: str) -> None:
    try:
        from docx.shared import Pt  # type: ignore
    except Exception:
        Pt = None  # type: ignore

    text = str(title or "").strip()
    if not text:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    if Pt is not None:
        run.font.size = Pt(14)
    try:
        run.font.name = "Arial"
    except Exception:
        pass
    # Titles: navy text (no background fill).
    _style_paragraph_text(paragraph, bold=True, color_hex=BRAND_NAVY_HEX)
    try:
        paragraph.paragraph_format.space_before = 12
        paragraph.paragraph_format.space_after = 4
    except Exception:
        pass


def _add_section_gap(doc, points: int = 12) -> None:
    paragraph = doc.add_paragraph("")
    try:
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = points
    except Exception:
        pass


def _style_table_header_row(table) -> None:
    try:
        from docx.shared import RGBColor  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        return

    if not table.rows:
        return

    def _shade_cell(cell, fill_hex: str) -> None:
        fill = (fill_hex or "").strip().lstrip("#").upper()
        if not fill:
            return
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    header_row = table.rows[0]
    for cell in header_row.cells:
        _shade_cell(cell, BRAND_NAVY_HEX)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                try:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                except Exception:
                    pass


def _load_draft(draft_id: str) -> Optional[Dict[str, Any]]:
    doc = FIRESTORE_CLIENT.collection(USER_STORY_DOCUMENT_DRAFTS_COLLECTION).document(draft_id).get()
    return doc.to_dict() if doc.exists else None


def _save_draft(draft_id: str, payload: Dict[str, Any]) -> None:
    FIRESTORE_CLIENT.collection(USER_STORY_DOCUMENT_DRAFTS_COLLECTION).document(draft_id).set(payload, merge=True)


def _normalize_document_payload(
    raw_document: Any,
    *,
    story: Dict[str, Any],
    epic: Dict[str, Any],
    project: Dict[str, Any],
) -> Dict[str, Any]:
    document: Dict[str, Any] = raw_document if isinstance(raw_document, dict) else {}

    story_statement = str(story.get("user_story") or story.get("title") or "").strip()
    story_id = str(story.get("user_story_id") or "").strip()
    dependencies = story.get("dependencies") or []
    if not isinstance(dependencies, list):
        dependencies = []

    def _stringify_bullets(value: Any) -> str:
        if isinstance(value, list):
            items = [str(item).strip() for item in value if str(item).strip()]
            if not items:
                return ""
            return "\n".join(f"- {item}" for item in items)
        return str(value or "").strip()

    document.setdefault("jira_id", story_id)
    document.setdefault("user_story_name", story_statement or story_id)
    document.setdefault("project_associated", str(project.get("name") or project.get("project_name") or "").strip())
    document.setdefault("responsible", str(story.get("assignee") or "").strip())
    document.setdefault("date", _now_display_date())
    document.setdefault("dependencies", ", ".join([str(d).strip() for d in dependencies if str(d).strip()]))

    # Provide a sensible default for description/scope if user story has a description but no document content yet.
    if _is_blank(document.get("description_and_scope")):
        fallback_description = str(story.get("description") or "").strip()
        if fallback_description:
            document["description_and_scope"] = fallback_description

    # Prefill sections that already exist on the user story so we don't ask redundant questions.
    if _is_blank(document.get("out_of_scope")):
        out_of_scope_text = _stringify_bullets(story.get("outOfScope") or story.get("out_of_scope"))
        if out_of_scope_text:
            document["out_of_scope"] = out_of_scope_text

    if _is_blank(document.get("acceptance_criteria")):
        acceptance_text = _stringify_bullets(story.get("acceptanceCriteria") or story.get("acceptance_criteria"))
        if acceptance_text:
            document["acceptance_criteria"] = acceptance_text

    if not isinstance(document.get(WIREFRAME_IMAGES_KEY), list):
        document[WIREFRAME_IMAGES_KEY] = []

    return document


def _missing_document_keys(document: Dict[str, Any]) -> List[str]:
    missing: List[str] = []
    for key in DOCUMENT_SECTION_KEYS:
        if key == "wireframe_mockup":
            images = document.get(WIREFRAME_IMAGES_KEY)
            has_images = False
            if isinstance(images, list):
                for item in images:
                    if isinstance(item, dict):
                        candidate = str(
                            item.get("path")
                            or item.get("url")
                            or item.get("inline_base64")
                            or item.get("data_base64")
                            or ""
                        ).strip()
                        if not candidate:
                            raw_bytes = item.get("inline_bytes")
                            if isinstance(raw_bytes, (bytes, bytearray)) and raw_bytes:
                                candidate = "__inline_bytes__"
                    else:
                        candidate = str(item or "").strip()
                    if candidate:
                        has_images = True
                        break
            if _is_blank(document.get(key)) and not has_images:
                missing.append(key)
            continue

        if _is_blank(document.get(key)):
            missing.append(key)
    return missing


async def start_user_story_document_draft(
    *,
    user_data: UserData,
    story_id: str,
) -> ResponseModel:
    story_response = get_user_story_by_id(
        story_id,
        user_data.get_user_id(),
        allow_member=True,
        user_email=user_data.get_email(),
    )
    if not story_response.success:
        return ResponseModel(success=False, message=story_response.message, data=None)

    story = story_response.data or {}
    epic_id = str(story.get("epic_id") or "").strip()
    if not epic_id:
        return ResponseModel(success=False, message="User story is missing epic_id", data=None)

    epic_response = get_epic_by_id(epic_id)
    if not epic_response.success:
        return ResponseModel(success=False, message="Epic not found", data=None)

    project_id = str(epic_response.data.get("project_id") or "").strip()
    project_response = get_project_by_id(
        project_id,
        user_data.get_user_id(),
        allow_member=True,
        user_email=user_data.get_email(),
    )
    project = project_response.data if project_response and project_response.success else {}

    document = _normalize_document_payload(
        story.get("document"),
        story=story,
        epic=epic_response.data,
        project=project or {},
    )

    missing = _missing_document_keys(document)
    questions = [{"key": key, "question": _get_question_text(key)} for key in missing[:MAX_DOCUMENT_QUESTIONS]]
    questions_keys = [q["key"] for q in questions]

    draft_id = f"usdoc_{_random_suffix(10)}"
    status = "questions" if questions else "ready"
    _save_draft(
        draft_id,
        {
            "draft_id": draft_id,
            "story_id": story_id,
            "epic_id": epic_id,
            "project_id": project_id,
            "status": status,
            "questions": questions,
            "missing_keys": questions_keys,
            "document": document,
            "created_at": _now_iso(),
            "updated_at": _now_iso(),
        },
    )

    return ResponseModel(
        success=True,
        message="Document draft started",
        data={"draft_id": draft_id, "status": status, "questions": questions},
    )


async def submit_user_story_document_answers(
    *,
    user_data: UserData,
    draft_id: str,
    answers: List[Dict[str, str]],
) -> ResponseModel:
    draft = _load_draft(draft_id)
    if not draft:
        return ResponseModel(success=False, message="Draft not found", data=None)

    story_id = str(draft.get("story_id") or "").strip()
    if not story_id:
        return ResponseModel(success=False, message="Draft is missing story_id", data=None)

    # Verify user still has access to the story.
    story_response = get_user_story_by_id(
        story_id,
        user_data.get_user_id(),
        allow_member=True,
        user_email=user_data.get_email(),
    )
    if not story_response.success:
        return ResponseModel(success=False, message=story_response.message, data=None)

    epic_id = str(draft.get("epic_id") or "").strip()
    epic_response = get_epic_by_id(epic_id) if epic_id else None
    project_id = str((epic_response.data or {}).get("project_id") or "") if epic_response and epic_response.success else ""
    project_response = (
        get_project_by_id(
            project_id,
            user_data.get_user_id(),
            allow_member=True,
            user_email=user_data.get_email(),
        )
        if project_id
        else None
    )
    project = project_response.data if project_response and project_response.success else {}

    document = _normalize_document_payload(
        draft.get("document") or {},
        story=story_response.data or {},
        epic=epic_response.data if epic_response and epic_response.success else {},
        project=project or {},
    )

    for item in answers or []:
        key = str(item.get("key") or "").strip()
        answer = str(item.get("answer") or "").strip()
        if not key:
            continue
        if answer:
            document[key] = answer

    # One clarification round: after answers, mark ready even if some fields remain empty.
    missing = _missing_document_keys(document)
    for key in missing:
        document[key] = document.get(key) or "N/A"

    _save_draft(
        draft_id,
        {
            "status": "ready",
            "questions": [],
            "missing_keys": [],
            "document": document,
            "updated_at": _now_iso(),
        },
    )

    return ResponseModel(
        success=True,
        message="Document ready",
        data={"draft_id": draft_id, "status": "ready", "questions": []},
    )


async def upload_user_story_document_wireframe_images(
    *,
    user_data: UserData,
    draft_id: str,
    files: List[UploadFile],
) -> ResponseModel:
    draft = _load_draft(draft_id)
    if not draft:
        return ResponseModel(success=False, message="Draft not found", data=None)

    story_id = str(draft.get("story_id") or "").strip()
    if not story_id:
        return ResponseModel(success=False, message="Draft is missing story_id", data=None)

    story_response = get_user_story_by_id(
        story_id,
        user_data.get_user_id(),
        allow_member=True,
        user_email=user_data.get_email(),
    )
    if not story_response.success:
        return ResponseModel(success=False, message=story_response.message, data=None)

    bucket = FIREBASE.storage_client
    using_inline_storage = bucket is None
    max_images = MAX_INLINE_WIREFRAME_IMAGES if using_inline_storage else MAX_WIREFRAME_IMAGES

    document = draft.get("document") if isinstance(draft.get("document"), dict) else {}
    existing = document.get(WIREFRAME_IMAGES_KEY)
    if not isinstance(existing, list):
        existing = []

    if len(existing) >= max_images:
        return ResponseModel(success=False, message="Maximum number of wireframe images reached", data=None)

    pending_files = [f for f in (files or []) if f is not None]
    if not pending_files:
        return ResponseModel(success=False, message="No files provided", data=None)

    if len(existing) + len(pending_files) > max_images:
        return ResponseModel(
            success=False,
            message=f"Too many images. Max allowed is {max_images}.",
            data=None,
        )

    uploaded: List[Dict[str, str]] = []
    for file in pending_files:
        content_type = str(file.content_type or "").strip().lower()
        if not content_type.startswith("image/"):
            return ResponseModel(
                success=False,
                message=f"Invalid file type for {file.filename or 'file'}. Only images are allowed.",
                data=None,
            )

        data = await file.read()
        if not data:
            continue
        max_bytes = MAX_INLINE_WIREFRAME_IMAGE_BYTES if using_inline_storage else MAX_WIREFRAME_IMAGE_BYTES
        if len(data) > max_bytes:
            return ResponseModel(
                success=False,
                message=(
                    f"File {file.filename or 'file'} is too large. "
                    f"Max allowed is {max_bytes // 1024}KB when Firebase Storage is not configured."
                    if using_inline_storage
                    else f"File {file.filename or 'file'} exceeds the {MAX_WIREFRAME_IMAGE_BYTES // (1024 * 1024)}MB limit."
                ),
                data=None,
            )

        raw_name = str(file.filename or "").strip() or f"wireframe_{_random_suffix(6)}.png"
        safe_name = "".join(ch if ch.isalnum() or ch in {"-", "_", "."} else "_" for ch in raw_name)
        safe_name = safe_name.strip("._") or f"wireframe_{_random_suffix(6)}.png"

        if using_inline_storage:
            uploaded.append(
                {
                    "inline_base64": base64.b64encode(data).decode("ascii"),
                    "filename": safe_name,
                    "content_type": content_type,
                }
            )
        else:
            storage_path = f"user-story-documents/{draft_id}/wireframes/{_random_suffix(6)}_{safe_name}"
            blob = bucket.blob(storage_path)
            blob.upload_from_string(data, content_type=content_type)

            url = ""
            try:
                blob.make_public()
                url = str(blob.public_url or "")
            except Exception:
                try:
                    url = str(blob.generate_signed_url(version="v4", expiration=3600) or "")
                except Exception:
                    url = ""

            uploaded.append(
                {
                    "path": storage_path,
                    "url": url,
                    "filename": safe_name,
                    "content_type": content_type,
                }
            )

    if not uploaded:
        return ResponseModel(success=False, message="No valid images uploaded", data=None)

    document = dict(document)
    document[WIREFRAME_IMAGES_KEY] = list(existing) + uploaded
    _save_draft(
        draft_id,
        {
            "document": document,
            "updated_at": _now_iso(),
        },
    )

    response_images: List[Dict[str, str]] = []
    for item in uploaded:
        if not isinstance(item, dict):
            continue
        response_images.append(
            {
                "path": str(item.get("path") or ""),
                "url": str(item.get("url") or ""),
                "filename": str(item.get("filename") or ""),
                "content_type": str(item.get("content_type") or ""),
                "stored_inline": "inline_base64" in item,
            }
        )

    return ResponseModel(
        success=True,
        message="Wireframe images uploaded",
        data={"draft_id": draft_id, "images": response_images},
    )


def _parse_user_story_statement(statement: str) -> Tuple[str, str, str]:
    raw = (statement or "").strip()
    if not raw:
        return "", "", ""

    normalized = re.sub(r"\s+", " ", raw).strip()
    lower = normalized.casefold()

    role_markers = ["as a ", "as an ", "as the ", "como "]
    need_markers = [" i want to ", " i want ", " i need to ", " i need ", " quiero ", " necesito "]
    value_markers = [" so that ", " in order to ", " para ", " con el fin de "]

    role_index = -1
    role_marker = ""
    for marker in role_markers:
        idx = lower.find(marker)
        if idx >= 0:
            role_index = idx
            role_marker = marker
            break

    if role_index < 0:
        return "", normalized, ""

    after_role = normalized[role_index + len(role_marker) :]
    after_role_lower = after_role.casefold()

    need_index = -1
    need_marker = ""
    for marker in need_markers:
        idx = after_role_lower.find(marker)
        if idx >= 0:
            need_index = idx
            need_marker = marker
            break

    if need_index < 0:
        return after_role.strip(" ,.-\n\t"), normalized, ""

    persona = after_role[:need_index].strip(" ,.-\n\t")
    rest = after_role[need_index + len(need_marker) :].strip()
    rest_lower = rest.casefold()

    value_index = -1
    value_marker = ""
    for marker in value_markers:
        idx = rest_lower.find(marker)
        if idx >= 0:
            value_index = idx
            value_marker = marker
            break

    if value_index >= 0:
        want_part = rest[:value_index].strip(" ,.-\n\t")
        so_that = rest[value_index + len(value_marker) :].strip(" ,.-\n\t")
        return persona, want_part, so_that

    return persona, rest.strip(" ,.-\n\t"), ""


def _add_section(doc, title: str, content: str) -> None:
    _add_section_header(doc, title)
    text = (content or "").strip()
    if not text:
        doc.add_paragraph("N/A")
        _add_section_gap(doc)
        return

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if len(lines) > 1:
        for ln in lines:
            doc.add_paragraph(ln, style="List Bullet")
        _add_section_gap(doc)
        return
    doc.add_paragraph(text)
    _add_section_gap(doc)


def _add_wireframe_section(doc, *, content: str, images: Any) -> None:
    _add_section_header(doc, "Wireframe / Mockup*")

    text = (content or "").strip()
    image_items = images if isinstance(images, list) else []
    has_images = False
    for item in image_items:
        if isinstance(item, dict):
            candidate = str(
                item.get("path")
                or item.get("url")
                or item.get("inline_base64")
                or item.get("data_base64")
                or ""
            ).strip()
            if not candidate:
                raw_bytes = item.get("inline_bytes")
                if isinstance(raw_bytes, (bytes, bytearray)) and raw_bytes:
                    candidate = "__inline_bytes__"
        else:
            candidate = str(item or "").strip()
        if candidate:
            has_images = True
            break

    if text:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        if len(lines) > 1:
            for ln in lines:
                doc.add_paragraph(ln, style="List Bullet")
        else:
            doc.add_paragraph(text)
    else:
        doc.add_paragraph("See images below." if has_images else "N/A")

    if not has_images:
        _add_section_gap(doc)
        return

    try:
        from docx.shared import Inches  # type: ignore
    except Exception:
        Inches = None  # type: ignore

    bucket = FIREBASE.storage_client
    for item in image_items:
        path = ""
        url = ""
        filename = ""
        inline_b64 = ""
        inline_bytes: Optional[bytes] = None
        if isinstance(item, dict):
            path = str(item.get("path") or "").strip()
            url = str(item.get("url") or "").strip()
            filename = str(item.get("filename") or "").strip()
            inline_b64 = str(item.get("inline_base64") or item.get("data_base64") or "").strip()
            raw_bytes = item.get("inline_bytes")
            if isinstance(raw_bytes, (bytes, bytearray)):
                inline_bytes = bytes(raw_bytes)
        else:
            path = str(item or "").strip()

        image_bytes: Optional[bytes] = None
        if inline_bytes:
            image_bytes = inline_bytes
        if inline_b64:
            try:
                image_bytes = base64.b64decode(inline_b64)
            except Exception:
                image_bytes = None
        if bucket and path:
            try:
                image_bytes = bucket.blob(path).download_as_bytes()
            except Exception:
                image_bytes = None

        if image_bytes:
            try:
                stream = BytesIO(image_bytes)
                if Inches is not None:
                    doc.add_picture(stream, width=Inches(6))
                else:
                    doc.add_picture(stream)
                if filename:
                    doc.add_paragraph(filename)
            except Exception:
                # Some PNG variants (e.g. unusual color profiles) can fail in python-docx.
                # Try re-encoding via Pillow and embed again.
                try:
                    from PIL import Image  # type: ignore

                    with Image.open(BytesIO(image_bytes)) as img:
                        # Ensure a standard pixel format.
                        if img.mode not in {"RGB", "RGBA"}:
                            img = img.convert("RGBA" if "A" in img.getbands() else "RGB")

                        converted = BytesIO()
                        img.save(converted, format="PNG")
                        converted.seek(0)
                        if Inches is not None:
                            doc.add_picture(converted, width=Inches(6))
                        else:
                            doc.add_picture(converted)
                        if filename:
                            doc.add_paragraph(filename)
                except Exception:
                    # Final fallback: at least include the filename so the user knows what was intended.
                    if filename:
                        doc.add_paragraph(f"[Image could not be embedded: {filename}]")
                continue
        elif url:
            doc.add_paragraph(url)

    _add_section_gap(doc)


def _set_table_style(table) -> None:
    try:
        table.style = "Table Grid"
    except Exception:
        return


def _parse_table_rows(content: Any, expected_columns: int) -> List[List[str]]:
    text = str(content or "").strip()
    if not text:
        return []

    rows: List[List[str]] = []
    for raw_line in text.splitlines():
        line = str(raw_line or "").strip()
        if not line:
            continue
        for prefix in ("- ", "* ", "• "):
            if line.startswith(prefix):
                line = line[len(prefix) :].strip()
                break
        if not line:
            continue

        cells = [part.strip() for part in line.split("|")]
        if len(cells) != expected_columns:
            rows.append([line] + [""] * max(0, expected_columns - 1))
            continue
        rows.append(cells)
    return rows


def _add_structured_table_rows(table, content: Any, expected_columns: int, start_row_index: int = 2) -> None:
    rows = _parse_table_rows(content, expected_columns)
    if not rows:
        rows = [[str(content or "N/A").strip() or "N/A"] + [""] * max(0, expected_columns - 1)]

    existing_rows = len(table.rows)
    while existing_rows < start_row_index + len(rows):
        table.add_row()
        existing_rows += 1

    for row_offset, row_values in enumerate(rows):
        row_index = start_row_index + row_offset
        for col_index in range(expected_columns):
            table.cell(row_index, col_index).text = row_values[col_index] if col_index < len(row_values) else ""


def _add_label_value_paragraph(doc, label: str, value: str) -> None:
    """
    Add a paragraph like `Label: value` where the `Label:` part is bold.
    """
    label_text = str(label or "").strip().rstrip(":")
    value_text = str(value or "").strip()
    if not label_text and not value_text:
        return

    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label_text}:")
    label_run.bold = True
    if value_text:
        paragraph.add_run(f" {value_text}")


def build_user_story_document_bytes(
    *,
    story: Dict[str, Any],
    project: Dict[str, Any],
    document: Dict[str, Any],
) -> Tuple[bytes, str]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("python-docx is required to generate DOCX files.") from exc

    doc = Document()
    _apply_document_defaults(doc)
    header_inserted = _add_header_image(doc)
    # If the Word header image is missing/unavailable (or failed to embed), add a visible top-of-body banner instead.
    if not header_inserted:
        _add_report_header(doc)

    _add_section_header(doc, "User Story / Requirement")
    _add_label_value_paragraph(doc, "Jira ID", str(document.get("jira_id") or story.get("user_story_id") or ""))
    _add_label_value_paragraph(
        doc,
        "User Story Name / Requirement",
        str(document.get("user_story_name") or story.get("user_story") or ""),
    )
    _add_label_value_paragraph(doc, "Project Associated", str(document.get("project_associated") or project.get("name") or ""))
    _add_label_value_paragraph(doc, "Responsible", str(document.get("responsible") or story.get("assignee") or ""))
    _add_label_value_paragraph(doc, "Date", str(document.get("date") or _now_display_date()))

    persona, want, so_that = _parse_user_story_statement(str(story.get("user_story") or ""))
    fallback_description = str(document.get("description_and_scope") or story.get("description") or "").strip()
    story_table = doc.add_table(rows=3, cols=2)
    _set_table_style(story_table)
    story_table.cell(0, 0).text = "As a …"
    story_table.cell(0, 1).text = persona or str(document.get("persona") or "").strip() or "N/A"
    story_table.cell(1, 0).text = "I want to …"
    story_table.cell(1, 1).text = want or fallback_description or str(story.get("user_story") or "").strip() or "N/A"
    story_table.cell(2, 0).text = "In order to …"
    story_table.cell(2, 1).text = so_that or str(document.get("in_order_to") or "").strip() or str(document.get("benefits") or "").strip() or "N/A"

    _add_section_gap(doc)
    _add_section(doc, "Description and Scope", str(document.get("description_and_scope") or ""))
    _add_section(doc, "Out of Scope", str(document.get("out_of_scope") or ""))
    _add_section(doc, "Preconditions", str(document.get("preconditions") or ""))
    _add_section(doc, "Entry Points", str(document.get("entry_points") or "N/A"))
    _add_section(doc, "Output Points*", str(document.get("output_points") or "N/A"))
    _add_section(doc, "Success Flow", str(document.get("success_flow") or ""))
    _add_wireframe_section(
        doc,
        content=str(document.get("wireframe_mockup") or ""),
        images=document.get(WIREFRAME_IMAGES_KEY),
    )

    _add_section_header(doc, "Field Description")
    doc.add_paragraph(
        "List and describe each data field involved, including business rules and their expected behavior."
    )
    field_table = doc.add_table(rows=2, cols=8)
    _set_table_style(field_table)
    field_headers = [
        "Element Name",
        "Data Name on the System",
        "Data - Source System",
        "Behavior",
        "Format",
        "Data Type",
        "Example",
        "Visibility when empty",
    ]
    for idx, header in enumerate(field_headers):
        field_table.cell(0, idx).text = header
    _style_table_header_row(field_table)
    placeholders = [
        "<Element>",
        "<Name of the element within the source system>",
        "<Name of the source system of the element>",
        "<Current behavior of the element>",
        "<CSV, TXT, HTML,JSON, etc.>",
        "<Numeric, alphanumeric, string, etc>",
        "<Example of the element>",
        "<YES / NO>",
    ]
    for idx, value in enumerate(placeholders):
        field_table.cell(1, idx).text = value
    _add_structured_table_rows(field_table, document.get("field_description"), 8, start_row_index=2)

    _add_section_gap(doc)
    _add_section_header(doc, "API Description*")
    api_table = doc.add_table(rows=2, cols=6)
    _set_table_style(api_table)
    api_headers = [
        "Source System",
        "Target System",
        "Connection Type",
        "Data Format",
        "Technical Viability",
        "Comments",
    ]
    for idx, header in enumerate(api_headers):
        api_table.cell(0, idx).text = header
    _style_table_header_row(api_table)
    api_placeholders = [
        "<System>",
        "<System>",
        "<Connection Type>",
        "<Data Format>",
        "<Compatible / Modification required>",
        "<Comments>",
    ]
    for idx, value in enumerate(api_placeholders):
        api_table.cell(1, idx).text = value
    _add_structured_table_rows(api_table, document.get("api_description"), 6, start_row_index=2)

    _add_section_gap(doc)
    _add_section(doc, "Acceptance Criteria", str(document.get("acceptance_criteria") or ""))
    _add_section(doc, "Test Scenarios", str(document.get("test_scenarios") or ""))
    _add_section(doc, "Dependencies", str(document.get("dependencies") or "N/A"))
    _add_section(doc, "Benefits", str(document.get("benefits") or "N/A"))

    _add_section_header(doc, "Estimation Dev.")
    doc.add_paragraph(
        "Time estimated for the development of this User Story. *This section must be filled by the development team."
    )
    estimation_table = doc.add_table(rows=3, cols=4)
    _set_table_style(estimation_table)
    for idx, header in enumerate(["Complexity Size", "Time Estimated", "Estimated by", "Estimation Date"]):
        estimation_table.cell(0, idx).text = header
    _style_table_header_row(estimation_table)
    estimation_table.cell(1, 0).text = "<Size>"
    estimation_table.cell(1, 1).text = "<Time>"
    estimation_table.cell(1, 2).text = "<Name>"
    estimation_table.cell(1, 3).text = "DD.MM.YYYY"
    estimation_table.cell(2, 0).text = str(story.get("tshirt_size") or story.get("tshirtSize") or "")
    estimation_table.cell(2, 1).text = str(document.get("estimation_dev") or "")
    estimation_table.cell(2, 2).text = str(document.get("estimated_by") or "")
    estimation_table.cell(2, 3).text = _now_display_date().replace(".", "/")

    _add_section_gap(doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_id = str(story.get("user_story_id") or story.get("id") or "user_story").strip().replace(" ", "_")
    filename = f"{safe_id}_User_Story.docx"
    return buffer.getvalue(), filename


async def build_user_story_document_download_response(
    *,
    user_data: UserData,
    draft_id: str,
) -> StreamingResponse:
    draft = _load_draft(draft_id)
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    story_id = str(draft.get("story_id") or "").strip()
    if not story_id:
        raise HTTPException(status_code=400, detail="Draft is missing story_id")

    story_response = get_user_story_by_id(
        story_id,
        user_data.get_user_id(),
        allow_member=True,
        user_email=user_data.get_email(),
    )
    if not story_response.success:
        status_code = 404 if "not found" in story_response.message.lower() else 403
        raise HTTPException(status_code=status_code, detail=story_response.message)

    epic_id = str(draft.get("epic_id") or "").strip()
    epic_response = get_epic_by_id(epic_id) if epic_id else None
    project_id = str((epic_response.data or {}).get("project_id") or "") if epic_response and epic_response.success else ""
    project_response = (
        get_project_by_id(
            project_id,
            user_data.get_user_id(),
            allow_member=True,
            user_email=user_data.get_email(),
        )
        if project_id
        else None
    )
    project = project_response.data if project_response and project_response.success else {}

    document = draft.get("document") if isinstance(draft.get("document"), dict) else {}
    normalized_document = _normalize_document_payload(
        document,
        story=story_response.data or {},
        epic=epic_response.data if epic_response and epic_response.success else {},
        project=project or {},
    )

    normalized_document = await _formalize_docx_document_fields(user_data=user_data, document=normalized_document)

    doc_bytes, filename = build_user_story_document_bytes(
        story=story_response.data or {},
        project=project or {},
        document=normalized_document,
    )
    encoded = quote(filename)
    return StreamingResponse(
        BytesIO(doc_bytes),
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )


def build_user_story_document_format_test_bytes() -> Tuple[bytes, str]:
    """
    Generate a fast, self-contained DOCX for quickly validating header/theme formatting.

    This intentionally avoids Firestore lookups and does not require a user story draft.
    """
    story = {
        "id": "docx_format_test",
        "user_story_id": "DOCX_TEST",
        "user_story": "As a Product Owner, I want a branded user story document so that I can review requirements consistently.",
        "description": "This is a test document generated to validate layout, colors, and header branding.",
        "dependencies": ["AUTH_LOGIN", "PROJECT_SETUP"],
        "storyPoints": 3,
        "effortHours": 2,
        "assignee": "Format Test",
    }
    project = {"name": "DOCX Format Test"}
    document = {
        "jira_id": "DOCX-TEST",
        "user_story_name": "DOCX Format Test Document",
        "project_associated": "DOCX Format Test",
        "responsible": "Format Test",
        "date": _now_display_date(),
        "description_and_scope": "- Validate header image\n- Validate navy/white section headers\n- Validate table header styling",
        "out_of_scope": "- Business logic correctness\n- Integration validations",
        "preconditions": "- None",
        "entry_points": "- Project > Epics > User Story Details",
        "output_points": "- A downloadable DOCX file with correct branding",
        "success_flow": "1) Click Test Document\n2) Download opens\n3) Header and section colors match",
        "wireframe_mockup": "N/A",
        "field_description": "N/A",
        "api_description": "N/A",
        "acceptance_criteria": "- Header is visible at top of page\n- Section headers use navy background\n- Table headers are navy with white text",
        "test_scenarios": "- Open in Word\n- Open in Google Docs\n- Print preview",
        "benefits": "- Faster iteration on formatting",
        "estimation_dev": "N/A",
    }
    return build_user_story_document_bytes(story=story, project=project, document=document)


async def build_user_story_document_download_response_with_wireframes(
    *,
    user_data: UserData,
    draft_id: str,
    files: List[UploadFile],
) -> StreamingResponse:
    """
    Build a DOCX download response for a user story document draft, embedding wireframe images
    from the current request WITHOUT persisting them (no Firebase Storage required).
    """
    draft = _load_draft(draft_id)
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    story_id = str(draft.get("story_id") or "").strip()
    if not story_id:
        raise HTTPException(status_code=400, detail="Draft is missing story_id")

    story_response = get_user_story_by_id(
        story_id,
        user_data.get_user_id(),
        allow_member=True,
        user_email=user_data.get_email(),
    )
    if not story_response.success:
        status_code = 404 if "not found" in story_response.message.lower() else 403
        raise HTTPException(status_code=status_code, detail=story_response.message)

    epic_id = str(draft.get("epic_id") or "").strip()
    epic_response = get_epic_by_id(epic_id) if epic_id else None
    project_id = str((epic_response.data or {}).get("project_id") or "") if epic_response and epic_response.success else ""
    project_response = (
        get_project_by_id(
            project_id,
            user_data.get_user_id(),
            allow_member=True,
            user_email=user_data.get_email(),
        )
        if project_id
        else None
    )
    project = project_response.data if project_response and project_response.success else {}

    document = draft.get("document") if isinstance(draft.get("document"), dict) else {}
    normalized_document = _normalize_document_payload(
        document,
        story=story_response.data or {},
        epic=epic_response.data if epic_response and epic_response.success else {},
        project=project or {},
    )

    pending_files = [f for f in (files or []) if f is not None]
    if len(pending_files) > MAX_WIREFRAME_IMAGES:
        raise HTTPException(status_code=400, detail=f"Too many images. Max allowed is {MAX_WIREFRAME_IMAGES}.")

    inline_images: List[Dict[str, Any]] = []
    for file in pending_files:
        content_type = str(file.content_type or "").strip().lower()
        if not content_type.startswith("image/"):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type for {file.filename or 'file'}. Only images are allowed.",
            )
        data = BytesIO(await file.read()).getvalue()
        if not data:
            continue
        if len(data) > MAX_WIREFRAME_IMAGE_BYTES:
            raise HTTPException(
                status_code=400,
                detail=f"File {file.filename or 'file'} exceeds the {MAX_WIREFRAME_IMAGE_BYTES // (1024 * 1024)}MB limit.",
            )
        inline_images.append(
            {
                "inline_bytes": data,
                "filename": str(file.filename or "").strip(),
                "content_type": content_type,
            }
        )

    if inline_images:
        normalized_document = dict(normalized_document)
        normalized_document[WIREFRAME_IMAGES_KEY] = inline_images

    normalized_document = await _formalize_docx_document_fields(user_data=user_data, document=normalized_document)

    doc_bytes, filename = build_user_story_document_bytes(
        story=story_response.data or {},
        project=project or {},
        document=normalized_document,
    )
    encoded = quote(filename)
    return StreamingResponse(
        BytesIO(doc_bytes),
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )
